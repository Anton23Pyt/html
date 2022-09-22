# -*- coding: cp1251 -*-
import re
import docx
from selenium import webdriver
from webdriver_manager.chrome import ChromeDriverManager
import time as ti
from bs4 import BeautifulSoup

def check(p,driver ):
    for i in range(0, len(p)):
        g = driver.find_element_by_css_selector(".editor--raw div.notranslate")
        g.send_keys(p[i])
        ti.sleep(2)
        o = driver.find_element_by_xpath("/html/body/div[1]/div[5]/div[3]/div/div[1]/div[2]/div[3]/button").click()
        fut = driver.find_element_by_css_selector("h2#symbols_count")
        while fut.text == "-":
            fut = driver.find_element_by_css_selector("h2#symbols_count")
            ti.sleep(3)
        x = driver.find_element_by_css_selector(
            "div:nth-of-type(2) .editor-common__EditingContainer-sc-1kwcpr1-4 div.editor-common__EditTextContainer-sc-1kwcpr1-5")
        p[i] = x.text
        g2 = driver.find_element_by_css_selector("button.iFhJuI").click()
    return p

def check1(p,driver):
    for i in range(0, len(p)):
        if p[i].find("http") == -1:
            g = driver.find_element_by_css_selector(".editor--raw div.notranslate")
            g.send_keys([p[i]])
            ti.sleep(2)
            o = driver.find_element_by_xpath(
                "/html/body/div[1]/div[5]/div[3]/div/div[1]/div[2]/div[3]/button").click()
            fut = driver.find_element_by_css_selector("h2#symbols_count")
            while fut.text == "-":
                fut = driver.find_element_by_css_selector("h2#symbols_count")
                ti.sleep(3)
            x = driver.find_element_by_css_selector(
                "div:nth-of-type(2) .editor-common__EditingContainer-sc-1kwcpr1-4 div.editor-common__EditTextContainer-sc-1kwcpr1-5")
            p[i] = x.text
            g2 = driver.find_element_by_css_selector("button.iFhJuI").click()
    return p

def check2(f,driver):
    g = driver.find_element_by_css_selector(".editor--raw div.notranslate")
    g.send_keys(f)
    ti.sleep(2)
    o = driver.find_element_by_xpath("/html/body/div[1]/div[5]/div[3]/div/div[1]/div[2]/div[3]/button").click()
    fut=driver.find_element_by_css_selector("h2#symbols_count")
    while fut.text=="-":
        fut = driver.find_element_by_css_selector("h2#symbols_count")
        ti.sleep(3)
    x = driver.find_element_by_css_selector(
        "div:nth-of-type(2) .editor-common__EditingContainer-sc-1kwcpr1-4 div.editor-common__EditTextContainer-sc-1kwcpr1-5")
    g2 = driver.find_element_by_css_selector("button.iFhJuI").click()
    return x.text

def find(tag):
    cut = []
    for i in str(tag):
        cut.append(i)
        if i == ">":
            yi = 0
            yo = 0
            for i in range(0, len(cut)):
                if cut[i] == "<":
                    yi = yi + 1
                if cut[i] == ">":
                    yo = yo + 1
            if yo == yi:
                break
    h = '\ '
    h = h.replace(' ', '')
    for i in range(0, len(cut)):
        h = h + cut[i]
    h = h + '(.*?)\</p>'
    return h



path=input("Введите распложение файла:")
doc = docx.Document(path)
driver = webdriver.Chrome(ChromeDriverManager().install())
driver.get("https://retext.ai/")
g1 = driver.find_element_by_css_selector("button.pretty-mount").click()
ti.sleep(3)
gh = "alekcey.prozorov@gmail.com"
gy = "Ale-X=x!tk1"
g1 = driver.find_element_by_css_selector("input.ce07c2f58")
g1.send_keys(gh)
g1 = driver.find_element_by_css_selector("input.ccd26e912")
g1.send_keys(gy)
g1 = driver.find_element_by_css_selector("button.c295cd12c").click()
ti.sleep(3)
text = []
for paragraph in doc.paragraphs:
    text.append(paragraph.text)
c='\n'.join(text)
soup = BeautifulSoup(c,'html.parser')
tag = soup.find_all('p')
if len(tag)>0:
    for i in range(0, len(tag)):
        z = tag[i]
        z2 = tag[i].text
        if str(z).find("href") == -1 and str((z.text)).find(" ") != -1 and z2[len(z2) - 1].find(":") == -1 and str(z).find("http") == -1:
            f = tag[i].text
            tag[i].string = check2(f, driver)
        if str(z).find("href") == -1 and str((z.text)).find(" ") != -1 and z2[len(z2) - 1].find(":") != -1 and str(z).find("http") == -1:
            p = []
            f = tag[i].text
            for s in re.split(r'(?<=[.!?…]) ', f):
                p.append(s)
            p = check(p, driver)
            ty = 0
            k = p[0]
            while ty < len(p) - 1:
                ty = ty + 1
                k = k + p[ty]
            tag[i].string = k
        if (str(z).find("href") != -1 or str(z).find("http") != -1) and str(z).find(" ") != -1:
            try:
                h=find(tag[i])
                g = re.findall(h, str(tag[i]))
                p=[]
                for s in re.split(r'(?<=[.!?…]) ', g[0]):
                    p.append(s)
                p = check1(p, driver)
                k = p[0]
                ty = 0
                while ty < len(p) - 1:
                    ty = ty + 1
                    k = k + p[ty]
                tag[i].string = k
            except:
                oi=0
tag2 = soup.find_all('li')
if len(tag2)>0:
    for i in range(0, len(tag2)):
        z = tag2[i]
        if str(z).find("href") == -1 and str((z.text)).find(" ") != -1 and str(z).find("http") == -1:
            f = tag2[i].text
            tag2[i].string = check2(f, driver)
driver.close()
c1=str(soup)
c1=c1.replace("&lt;","<")
c1=c1.replace("&gt;",">")
doc2 = docx.Document()
par=doc2.add_paragraph(c1)
doc2.save(path)